from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent
SOURCE_MD = ROOT / "MSH3_COURSE_PAPER_DRAFT.md"
OUTPUT_DOCX = ROOT / "MSH3_COURSE_PAPER_FORMATTED.docx"
TEMPLATE_DOC = PROJECT_ROOT / "文献检索与论文写作结课论文模板-(1).doc"
SOFFICE = Path(os.environ.get("SOFFICE", r"C:\Program Files\LibreOffice\program\soffice.com"))

CN_BODY = "宋体"
CN_HEADING = "黑体"
EN_BODY = "Times New Roman"
EN_MONO = "Courier New"

INK = RGBColor(0, 0, 0)
MUTED = RGBColor(88, 88, 88)
ACCENT = RGBColor(31, 77, 120)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F5F7FA"
GRID = "B8C2CC"

TABLE_CAPTIONS = [
    "表 1 MSH3文献检索与筛选流程概要",
    "表 2 MSH3常用检测技术比较",
    "表 3 MSH3临床证据与案例矩阵",
]


@dataclass
class Element:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)
    code_lines: list[str] = field(default_factory=list)


def run(command: list[str | os.PathLike[str]], cwd: Path = ROOT) -> None:
    completed = subprocess.run([str(part) for part in command], cwd=cwd, check=False)
    if completed.returncode != 0:
        raise RuntimeError(f"command failed with exit code {completed.returncode}: {command}")


def convert_reference_doc(build_dir: Path) -> Path | None:
    if not TEMPLATE_DOC.exists() or not SOFFICE.exists():
        return None
    run([SOFFICE, "--headless", "--norestore", "--convert-to", "docx", "--outdir", build_dir, TEMPLATE_DOC])
    candidates = list(build_dir.glob("*.docx"))
    return candidates[0] if candidates else None


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(
    run,
    *,
    east_asia: str = CN_BODY,
    latin: str = EN_BODY,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, *, east_asia: str, latin: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, east_asia=CN_BODY, latin=EN_BODY, size=12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size, before, after in [
        ("Heading 1", 16, 12, 8),
        ("Heading 2", 14, 8, 4),
        ("Heading 3", 13, 6, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, east_asia=CN_HEADING, latin=EN_BODY, size=size, bold=False, color=INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Reference Item" not in doc.styles:
        ref = doc.styles.add_style("Reference Item", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["Reference Item"]
    set_style_font(ref, east_asia=CN_BODY, latin=EN_BODY, size=10.5)
    ref.paragraph_format.left_indent = Cm(0.74)
    ref.paragraph_format.first_line_indent = Cm(-0.74)
    ref.paragraph_format.line_spacing = 1.15
    ref.paragraph_format.space_after = Pt(4)

    if "Code Block" not in doc.styles:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    set_style_font(code, east_asia=CN_BODY, latin=EN_MONO, size=9.3)
    code.paragraph_format.line_spacing = 1.15
    code.paragraph_format.space_after = Pt(0)


def set_page_setup(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10.5)


def set_body_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("《基因检测与解读》课程论文")
    set_run_font(run, size=9.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = footer_p.add_run("- ")
    set_run_font(left, size=10.5, color=MUTED)
    add_page_field(footer_p)
    right = footer_p.add_run(" -")
    set_run_font(right, size=10.5, color=MUTED)


def set_table_width(table, widths: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths]
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, cell in enumerate(row.cells):
            if index >= len(widths_dxa):
                continue
            cell.width = Inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = GRID, size: str = "4", value: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 130, bottom: int = 100, end: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def format_table(table, widths: list[float], *, font_size: float = 9.4) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_width(table, widths)
    if table.rows:
        mark_header_row(table.rows[0])
    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_data_table(doc: Document, caption: str, rows: list[list[str]], widths: list[float]) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    format_table(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_flow_figure(doc: Document, caption: str, lines: list[str]) -> None:
    clean_lines = [line.strip() for line in lines if line.strip() and line.strip() != "↓"]
    p_break = doc.add_paragraph()
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after = Pt(0)
    p_break.add_run().add_break(WD_BREAK.PAGE)
    add_caption(doc, caption.replace("（文字版）", ""))
    rows = len(clean_lines) * 2 - 1
    table = doc.add_table(rows=rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.0])
    step_index = 0
    for row_index, row in enumerate(table.rows):
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=90, start=160, bottom=90, end=160)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if row_index % 2 == 0:
            set_cell_shading(cell, "F3F7FB")
            set_cell_borders(cell, color="AFC2D8")
            run = p.add_run(clean_lines[step_index])
            set_run_font(run, size=10)
            step_index += 1
        else:
            set_cell_borders(cell, color="FFFFFF", value="nil")
            run = p.add_run("↓")
            set_run_font(run, size=12, color=ACCENT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw_cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in raw_cells)
        if not is_separator:
            rows.append(raw_cells)
        i += 1
    return rows, i


def parse_markdown() -> tuple[str, str, list[Element]]:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = next(line.removeprefix("题目：").strip() for line in lines if line.startswith("题目："))
    english_title = next(line.removeprefix("英文题名：").strip() for line in lines if line.startswith("英文题名："))

    elements: list[Element] = []
    i = 0
    pending_figure = ""
    table_count = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped.startswith("# ") or stripped.startswith("题目：") or stripped.startswith("英文题名："):
            i += 1
            continue

        if stripped.startswith("## "):
            elements.append(Element("heading", stripped[3:].strip(), 1))
            i += 1
            continue

        if stripped.startswith("### "):
            elements.append(Element("heading", stripped[4:].strip(), 2))
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                caption = TABLE_CAPTIONS[table_count] if table_count < len(TABLE_CAPTIONS) else f"表 {table_count + 1}"
                table_count += 1
                elements.append(Element("table", caption, rows=rows))
            continue

        if stripped.startswith("图 1"):
            pending_figure = stripped
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if pending_figure:
                elements.append(Element("figure", pending_figure, code_lines=code_lines))
                pending_figure = ""
            else:
                elements.append(Element("code", code_lines=code_lines))
            continue

        elements.append(Element("paragraph", stripped))
        i += 1

    return title, english_title, elements


def collect_toc_items(elements: list[Element]) -> list[Element]:
    return [element for element in elements if element.kind == "heading" and element.level == 1]


def add_cover(doc: Document, title: str, english_title: str) -> None:
    section = doc.sections[0]
    set_page_setup(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

    for text, size, after in [
        ("《基因检测与解读》", 21, 6),
        ("结课论文", 23, 20),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        set_run_font(run, east_asia=CN_HEADING, size=size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    label = p.add_run("论文题目")
    set_run_font(label, east_asia=CN_HEADING, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, east_asia=CN_HEADING, size=17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(english_title)
    set_run_font(run, east_asia=CN_BODY, latin=EN_BODY, size=11, italic=True, color=MUTED)

    rows = [
        ("学院", "生命与健康管理学院"),
        ("专业", "医学检验技术"),
        ("班级", "25医检专升本"),
        ("课程", "基因检测与解读"),
        ("姓名", ""),
        ("学号", ""),
        ("提交日期", "2026年6月"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.45, 4.65])
    for row_index, (label_text, value) in enumerate(rows):
        left, right = table.rows[row_index].cells
        left.text = label_text
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell, color="C9D2DC")
            set_cell_margins(cell, top=95, start=150, bottom=95, end=150)
            if cell is left:
                set_cell_shading(cell, "F1F4F8")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell is left else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=11, bold=(cell is left))


def add_toc(doc: Document, toc_items: list[Element], page_numbers: dict[str, int] | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("目   录")
    set_run_font(run, east_asia=CN_HEADING, size=16)

    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.3 if re.match(r"^\d+\s", item.text) else 0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        title_run = p.add_run(item.text)
        set_run_font(title_run, size=12)
        p.add_run("\t")
        page_run = p.add_run("" if not page_numbers else str(page_numbers.get(item.text, "")))
        set_run_font(page_run, size=12)


def add_heading(doc: Document, text: str, level: int, *, page_break_before: bool = False) -> None:
    if page_break_before:
        p_break = doc.add_paragraph()
        p_break.paragraph_format.space_before = Pt(0)
        p_break.paragraph_format.space_after = Pt(0)
        p_break.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph(style="Heading 1" if level == 1 else "Heading 2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = None
    if level == 1 and not re.match(r"^\d+\s", text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, east_asia=CN_HEADING, latin=EN_BODY, size=16 if level == 1 else 14)


def add_body_paragraph(doc: Document, text: str, current_h1: str) -> None:
    if re.match(r"^\[\d+\]", text):
        p = doc.add_paragraph(style="Reference Item")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.replace("\u00a0", " "))
        set_run_font(run, size=10.5)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text.startswith("关键词") or text.startswith("Keywords"):
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.25
        label, sep, rest = text.partition("：")
        if not sep:
            label, sep, rest = text.partition(":")
        label_run = p.add_run(label + (sep or ""))
        set_run_font(label_run, size=12, bold=True)
        rest_run = p.add_run(rest)
        set_run_font(rest_run, size=11.5)
    elif current_h1 in {"摘要", "Abstract"}:
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text)
        set_run_font(run, size=11.5)
    else:
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        set_run_font(run, size=12)


def table_widths(rows: list[list[str]]) -> list[float]:
    cols = len(rows[0])
    if cols == 2:
        return [1.6, 5.2]
    if cols == 3:
        return [1.55, 3.25, 2.0]
    if cols == 4:
        return [1.45, 2.05, 1.65, 1.7]
    return [6.8 / cols] * cols


def add_body(doc: Document, elements: list[Element]) -> None:
    current_h1 = ""
    seen_body_heading = False
    for element in elements:
        if element.kind == "heading":
            if element.level == 1:
                current_h1 = element.text
                force_page = seen_body_heading and element.text in {"引言", "1 文献检索策略与筛选过程"}
                add_heading(doc, element.text, element.level, page_break_before=force_page)
                seen_body_heading = True
            else:
                add_heading(doc, element.text, element.level)
        elif element.kind == "table":
            add_data_table(doc, element.text, element.rows, table_widths(element.rows))
        elif element.kind == "figure":
            add_flow_figure(doc, element.text, element.code_lines)
        elif element.kind == "code":
            add_code_block(doc, element.code_lines)
        elif element.kind == "paragraph":
            add_body_paragraph(doc, element.text, current_h1)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_borders(cell, color="D9DEE7")
    set_cell_margins(cell)
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "Code Block"
        p.paragraph_format.first_line_indent = None
        run = p.add_run(line)
        set_run_font(run, east_asia=CN_BODY, latin=EN_MONO, size=9.2)


def scrub_core_properties(doc: Document, title: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = "基因检测与解读课程论文"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "MSH3; 错配修复; 二代测序; 腺瘤性息肉病; 结直肠癌"
    props.comments = ""
    props.category = ""


def build_document(reference_docx: Path | None, page_numbers: dict[str, int] | None = None, output: Path = OUTPUT_DOCX) -> list[Element]:
    title, english_title, elements = parse_markdown()
    doc = Document(reference_docx) if reference_docx else Document()
    clear_document(doc)
    configure_styles(doc)
    scrub_core_properties(doc, title)

    add_cover(doc, title, english_title)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_setup(toc_section)
    toc_section.header.is_linked_to_previous = False
    toc_section.footer.is_linked_to_previous = False
    toc_section.header.paragraphs[0].text = ""
    toc_section.footer.paragraphs[0].text = ""
    add_toc(doc, collect_toc_items(elements), page_numbers)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_setup(body_section)
    restart_page_numbering(body_section, 1)
    set_body_header_footer(body_section)
    add_body(doc, elements)

    doc.save(output)
    return elements


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    if not SOFFICE.exists():
        raise FileNotFoundError(SOFFICE)
    run([SOFFICE, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", output_dir, docx_path])
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)
    return pdf_path


def extract_toc_page_numbers(pdf_path: Path, toc_items: list[Element]) -> dict[str, int]:
    try:
        import pdfplumber
    except ImportError:
        return {}

    found: dict[str, int] = {}
    front_matter_pages = 2
    with pdfplumber.open(pdf_path) as pdf:
        for physical_page, page in enumerate(pdf.pages, start=1):
            if physical_page <= front_matter_pages:
                continue
            text = page.extract_text() or ""
            lines = {line.strip() for line in text.splitlines() if line.strip()}
            body_page = physical_page - front_matter_pages
            for item in toc_items:
                if item.text not in found and item.text in lines:
                    found[item.text] = body_page
    return found


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    with tempfile.TemporaryDirectory(prefix="msh3-course-paper-docx-") as temp_dir:
        build_dir = Path(temp_dir)
        reference_docx = convert_reference_doc(build_dir)
        elements = build_document(reference_docx, page_numbers=None, output=OUTPUT_DOCX)

        try:
            pdf_dir = build_dir / "pdf"
            pdf_dir.mkdir()
            pdf_path = convert_docx_to_pdf(OUTPUT_DOCX, pdf_dir)
            page_numbers = extract_toc_page_numbers(pdf_path, collect_toc_items(elements))
        except Exception as exc:
            print(f"warning: static TOC page numbers were not generated: {exc}")
            page_numbers = {}

        build_document(reference_docx, page_numbers=page_numbers, output=OUTPUT_DOCX)

    print(f"wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()

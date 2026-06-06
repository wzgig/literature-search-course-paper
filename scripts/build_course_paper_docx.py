from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "COURSE_PAPER_DRAFT.md"
OUTPUT_DOCX = ROOT / "COURSE_PAPER_FORMATTED.docx"
TEMPLATE_DOC = ROOT / "文献检索与论文写作结课论文模板-(1).doc"
SOFFICE = Path(os.environ.get("SOFFICE", r"C:\Program Files\LibreOffice\program\soffice.com"))

CN_BODY = "宋体"
CN_HEADING = "黑体"
EN_BODY = "Times New Roman"
EN_MONO = "Courier New"

BLUE = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"
GRID = "B8C2CC"


TABLES = {
    "表 1 文献检索与筛选流程概要": (
        [
            ["步骤", "核心信息"],
            ["PubMed 宽检索", "围绕 sFlt-1、PlGF、sFlt-1/PlGF ratio 与 preeclampsia 建立英文候选池，共 2305 条结果。"],
            ["系统综述检索", "增加 systematic review、meta-analysis 限定，共 47 条结果，优先筛选综合证据。"],
            ["临床应用检索", "增加 clinical utility、implementation、rule out 等关键词，共 135 条结果，筛选管理决策相关文献。"],
            ["指南定向检索", "定向纳入 NICE PLGF-based testing guidance、ISSHP 指南和临床解释综述。"],
            ["中文补充检索", "使用“子痫前期”“sFlt-1”“胎盘生长因子”“PlGF”等关键词，纳入胡吉霞等、杨岚等中文研究。"],
            ["最终筛选", "按指南、Meta 分析、多中心研究、方法学研究和中文临床研究分层，形成 18 条临时参考文献。"],
        ],
        [1.35, 5.15],
    ),
    "表 2 sFlt-1、PlGF 与 sFlt-1/PlGF 比值比较": (
        [
            ["指标", "核心解释", "临床检验定位"],
            ["sFlt-1", "抗血管生成相关因子，升高提示抗血管生成状态增强。", "用于机制解释和联合检测，但单项结果受孕周及个体差异影响。"],
            ["PlGF", "促血管生成因子，降低提示胎盘功能和血管生成支持不足。", "部分 PLGF-based tests 已进入指南推荐，可用于疑似早产期子痫前期辅助判断。"],
            ["sFlt-1/PlGF 比值", "同时整合 sFlt-1 升高和 PlGF 降低，反映血管生成失衡。", "适合短期排除、风险分层和辅助诊断；阳性结果不能单独确诊。"],
            ["多指标联合模型", "将血管生成因子与其他标志物或临床信息结合。", "可用于复杂场景和不良结局预测，但成本、解释和路径要求更高。"],
        ],
        [1.35, 2.75, 2.40],
    ),
    "表 3 核心证据矩阵": (
        [
            ["文献或来源", "证据重点", "本文用途"],
            ["NICE PLGF-based testing guidance[1]", "PLGF 或 sFlt-1/PlGF 检测应与标准临床评估联合使用。", "确定临床应用边界。"],
            ["Zeisler 等 PROGNOSIS 研究[2]", "比值 ≤38 对 1 周内无子痫前期 NPV 为 99.3%。", "短期排除价值核心证据。"],
            ["Zhang 等 2025 Meta 分析[3]", "sFlt-1/PlGF AUC 0.92，高于单项 sFlt-1 或 PlGF。", "总体预测性能和局限。"],
            ["Zhao 等 2017 Meta 分析[4]", "双截断策略支持低风险、高风险和灰区分层。", "截断值解释和风险分层。"],
            ["Espinoza 等 2024 研究[5]", "MoM 11.5 预测 2 周内重症特征，AUC 0.91。", "重症短期风险分层。"],
            ["胡吉霞等 2022 研究[6]", "比值预测早发型子痫前期 AUC 0.873。", "国内中文临床证据。"],
            ["杨岚等 2024 研究[12]", "PLGF、SFLT-1、GLYFN 三项联合 AUC 0.986。", "多标志物联合证据。"],
            ["Lafuente-Ganuza 等 2020 研究[13]", "sFlt-1/PlGF 联合 NT-proBNP 可改善 rule-in/rule-out。", "局限性和未来方向。"],
        ],
        [2.00, 3.10, 1.40],
    ),
}

FLOW_STEPS = [
    "疑似或高风险子痫前期孕妇",
    "核对孕周、单胎/多胎、症状、血压、尿蛋白和常规实验室指标",
    "采集血清/血浆并检测 sFlt-1、PlGF，计算 sFlt-1/PlGF 比值",
    "结合平台说明书、孕周范围和研究证据解释结果",
    "低风险：提示短期发生子痫前期可能性较低，继续临床随访",
    "灰区：结合临床表现、胎儿情况和动态复查判断",
    "高风险：提示风险升高，加强监测、转诊或住院评估",
    "检验科复核质控、报告备注并与临床沟通，避免将比值单独作为确诊依据",
]


@dataclass
class Element:
    kind: str
    text: str
    level: int = 0


def run(command: list[str | os.PathLike[str]], cwd: Path = ROOT) -> None:
    completed = subprocess.run([str(part) for part in command], cwd=cwd, check=False)
    if completed.returncode != 0:
        raise RuntimeError(f"command failed with exit code {completed.returncode}: {command}")


def convert_reference_doc(build_dir: Path) -> Path | None:
    if not TEMPLATE_DOC.exists() or not SOFFICE.exists():
        return None
    run([SOFFICE, "--headless", "--norestore", "--convert-to", "docx", "--outdir", build_dir, TEMPLATE_DOC])
    candidates = list(build_dir.glob("*.docx"))
    return candidates[0] if candidates else None


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(
    run,
    *,
    east_asia: str = CN_BODY,
    latin: str = EN_BODY,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, *, east_asia: str, latin: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, east_asia=CN_BODY, latin=EN_BODY, size=12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size, before, after in [
        ("Heading 1", 16, 12, 8),
        ("Heading 2", 14, 8, 4),
        ("Heading 3", 13, 6, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, east_asia=CN_HEADING, latin=EN_BODY, size=size, bold=False, color=INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Course Code" not in doc.styles:
        code = doc.styles.add_style("Course Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Course Code"]
    set_style_font(code, east_asia=CN_BODY, latin=EN_MONO, size=9.5)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.15

    if "Reference Item" not in doc.styles:
        ref = doc.styles.add_style("Reference Item", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["Reference Item"]
    set_style_font(ref, east_asia=CN_BODY, latin=EN_BODY, size=10.5)
    ref.paragraph_format.left_indent = Cm(0.74)
    ref.paragraph_format.first_line_indent = Cm(-0.74)
    ref.paragraph_format.line_spacing = 1.15
    ref.paragraph_format.space_after = Pt(4)


def set_page_setup(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10.5)


def set_body_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("《文献检索与论文写作》课程论文")
    set_run_font(run, size=9.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = footer_p.add_run("- ")
    set_run_font(left, size=10.5, color=MUTED)
    add_page_field(footer_p)
    right = footer_p.add_run(" -")
    set_run_font(right, size=10.5, color=MUTED)


def set_table_width(table, widths: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths]
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, cell in enumerate(row.cells):
            if index >= len(widths_dxa):
                continue
            cell.width = Inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = GRID, size: str = "4", value: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 130, bottom: int = 100, end: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_table(table, widths: list[float], *, header_fill: str = LIGHT_BLUE, font_size: float = 9.5) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_width(table, widths)
    if table.rows:
        mark_header_row(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=False)


def add_data_table(doc: Document, caption: str, data: list[list[str]], widths: list[float]) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    format_table(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_borders(cell, color="D9DEE7")
    set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
    cell.text = ""
    for idx, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "Course Code"
        p.paragraph_format.first_line_indent = None
        run = p.add_run(line)
        set_run_font(run, east_asia=CN_BODY, latin=EN_MONO, size=9.2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_flow_figure(doc: Document, caption: str) -> None:
    add_caption(doc, caption.replace("（文字版）", ""))
    rows = len(FLOW_STEPS) * 2 - 1
    table = doc.add_table(rows=rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5.95])
    step_index = 0
    for row_index, row in enumerate(table.rows):
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=90, start=160, bottom=90, end=160)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if row_index % 2 == 0:
            set_cell_shading(cell, "F3F7FB")
            set_cell_borders(cell, color="AFC2D8")
            run = p.add_run(FLOW_STEPS[step_index])
            set_run_font(run, size=10.5)
            step_index += 1
        else:
            set_cell_borders(cell, color="FFFFFF", value="nil")
            run = p.add_run("↓")
            set_run_font(run, size=12, color=BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_keyword_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    label, _, rest = text.partition("：")
    if not rest:
        label, _, rest = text.partition(":")
        sep = ": "
    else:
        sep = "："
    label_run = p.add_run(label + sep)
    set_run_font(label_run, size=12, bold=True)
    rest_run = p.add_run(rest)
    set_run_font(rest_run, size=12)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Reference Item")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.replace("\u00a0", " "))
    set_run_font(run, size=10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = None
    if level == 1 and (not re.match(r"^\d+\s", text)):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(10)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, east_asia=CN_HEADING, latin=EN_BODY, size=16 if level == 1 else 14, bold=False)


def parse_markdown() -> tuple[str, str, list[Element]]:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = next(line.removeprefix("题目：").strip() for line in lines if line.startswith("题目："))
    english_title = next(line.removeprefix("英文题名：").strip() for line in lines if line.startswith("英文题名："))

    elements: list[Element] = []
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped or stripped.startswith("# 课程论文初稿") or stripped.startswith("版本：") or stripped.startswith("题目：") or stripped.startswith("英文题名：") or stripped.startswith("> 说明"):
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            block = "\n".join(code_lines).strip()
            if "疑似或高风险子痫前期孕妇" not in block:
                elements.append(Element("code", block))
            continue

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading == "临时参考文献":
                heading = "参考文献"
            elements.append(Element("heading", heading, 1))
            i += 1
            continue

        if stripped.startswith("### "):
            elements.append(Element("heading", stripped[4:].strip(), 2))
            i += 1
            continue

        if stripped in TABLES:
            elements.append(Element("table", stripped))
            i += 1
            while i < len(lines) and (lines[i].strip().startswith("|") or not lines[i].strip()):
                i += 1
            continue

        if stripped.startswith("图 1"):
            elements.append(Element("figure", stripped))
            i += 1
            while i < len(lines):
                if lines[i].strip().startswith("## "):
                    break
                i += 1
            continue

        if stripped.startswith("|"):
            i += 1
            continue

        elements.append(Element("paragraph", stripped))
        i += 1

    return title, english_title, elements


def collect_toc_items(elements: list[Element]) -> list[Element]:
    return [element for element in elements if element.kind == "heading" and element.level == 1]


def add_cover(doc: Document, title: str, english_title: str) -> None:
    section = doc.sections[0]
    set_page_setup(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

    for text, size, after in [
        ("《文献检索与论文写作》", 21, 6),
        ("结课论文", 23, 18),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        set_run_font(run, east_asia=CN_HEADING, size=size, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    label = p.add_run("论文题目")
    set_run_font(label, east_asia=CN_HEADING, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(title)
    set_run_font(run, east_asia=CN_HEADING, size=16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(english_title)
    set_run_font(run, east_asia=CN_BODY, latin=EN_BODY, size=11, italic=True, color=MUTED)

    rows = [
        ("学院", "生命与健康管理学院"),
        ("专业", "医学检验技术"),
        ("课程", "文献检索与论文写作"),
        ("姓名", ""),
        ("学号", ""),
        ("提交日期", "2026年6月6日"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.45, 4.65])
    for row_index, (label_text, value) in enumerate(rows):
        left, right = table.rows[row_index].cells
        left.text = label_text
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell, color="C9D2DC")
            set_cell_margins(cell, top=95, start=150, bottom=95, end=150)
            if cell is left:
                set_cell_shading(cell, "F1F4F8")
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = None
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell is left else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, size=11, bold=(cell is left))


def add_toc(doc: Document, toc_items: list[Element], page_numbers: dict[str, int] | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("目    录")
    set_run_font(run, east_asia=CN_HEADING, size=16)

    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.4 if re.match(r"^\d+\s", item.text) else 0)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        title_run = p.add_run(item.text)
        set_run_font(title_run, size=12)
        p.add_run("\t")
        number = "" if not page_numbers else str(page_numbers.get(item.text, ""))
        page_run = p.add_run(number)
        set_run_font(page_run, size=12)


def add_body(doc: Document, elements: list[Element]) -> None:
    for element in elements:
        if element.kind == "heading":
            add_heading(doc, element.text, element.level)
        elif element.kind == "table":
            data, widths = TABLES[element.text]
            add_data_table(doc, element.text, data, widths)
        elif element.kind == "figure":
            add_flow_figure(doc, element.text)
        elif element.kind == "code":
            add_code_block(doc, element.text)
        elif element.kind == "paragraph":
            if re.match(r"^\[\d+\]", element.text):
                add_reference(doc, element.text)
            elif element.text.startswith("关键词") or element.text.startswith("Keywords"):
                add_keyword_paragraph(doc, element.text)
            else:
                add_body_paragraph(doc, element.text)


def scrub_core_properties(doc: Document, title: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = "文献检索与论文写作课程论文"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "课程论文; 文献检索; 子痫前期; sFlt-1/PlGF"
    props.comments = ""
    props.category = ""


def build_document(reference_docx: Path | None, page_numbers: dict[str, int] | None = None, output: Path = OUTPUT_DOCX) -> list[Element]:
    title, english_title, elements = parse_markdown()
    doc = Document(reference_docx) if reference_docx else Document()
    clear_document(doc)
    configure_styles(doc)
    scrub_core_properties(doc, title)

    add_cover(doc, title, english_title)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_setup(toc_section)
    toc_section.header.is_linked_to_previous = False
    toc_section.footer.is_linked_to_previous = False
    toc_section.header.paragraphs[0].text = ""
    toc_section.footer.paragraphs[0].text = ""
    add_toc(doc, collect_toc_items(elements), page_numbers)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_setup(body_section)
    restart_page_numbering(body_section, 1)
    set_body_header_footer(body_section)
    add_body(doc, elements)

    doc.save(output)
    return elements


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    if not SOFFICE.exists():
        raise FileNotFoundError(SOFFICE)
    run([SOFFICE, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", output_dir, docx_path])
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)
    return pdf_path


def extract_toc_page_numbers(pdf_path: Path, toc_items: list[Element]) -> dict[str, int]:
    try:
        import pdfplumber
    except ImportError:
        return {}

    targets = [item.text for item in toc_items]
    found: dict[str, int] = {}
    with pdfplumber.open(pdf_path) as pdf:
        first_body_page = 3
        for physical_page, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if "《文献检索与论文写作》课程论文" in text:
                first_body_page = physical_page
                break
        front_matter_pages = first_body_page - 1
        for physical_page, page in enumerate(pdf.pages, start=1):
            if physical_page <= front_matter_pages:
                continue
            text = page.extract_text() or ""
            lines = {line.strip() for line in text.splitlines() if line.strip()}
            body_page = physical_page - front_matter_pages
            for target in targets:
                if target not in found and target in lines:
                    found[target] = body_page
    return found


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    with tempfile.TemporaryDirectory(prefix="course-paper-docx-build-") as temp_dir:
        build_dir = Path(temp_dir)
        reference_docx = convert_reference_doc(build_dir)
        elements = build_document(reference_docx, page_numbers=None, output=OUTPUT_DOCX)

        try:
            pdf_dir = build_dir / "pdf"
            pdf_dir.mkdir()
            pdf_path = convert_docx_to_pdf(OUTPUT_DOCX, pdf_dir)
            page_numbers = extract_toc_page_numbers(pdf_path, collect_toc_items(elements))
        except Exception as exc:
            print(f"warning: static TOC page numbers were not generated: {exc}")
            page_numbers = {}

        build_document(reference_docx, page_numbers=page_numbers, output=OUTPUT_DOCX)

    print(f"wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
